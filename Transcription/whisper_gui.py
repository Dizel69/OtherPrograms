import os
import whisper
import ttkbootstrap as ttk  # Используем ttkbootstrap
from ttkbootstrap.constants import *  # Для управления стилями
from tkinter import filedialog, messagebox, Toplevel, Label
from threading import Thread
from docx import Document  # Для сохранения в .docx

# Глобальная модель Whisper
model = None
model_name = "large"  # Модель по умолчанию

# Список моделей с описанием
MODEL_OPTIONS = {
    "tiny": "Модель tiny: самая быстрая, но менее точная.",
    "base": "Модель base: базовая модель для быстрого перевода.",
    "small": "Модель small: компромисс между скоростью и качеством.",
    "medium": "Модель medium: точнее, но медленнее.",
    "large": "Модель large: самая точная, требует больше ресурсов.",
}

# Функция загрузки выбранной модели
def load_model():
    global model
    model = whisper.load_model(model_name)
    messagebox.showinfo("Готово", f"Модель '{model_name}' загружена успешно!")

# Функция транскрибации
def transcribe_audio(file_path):
    try:
        if model is None:
            raise RuntimeError("Модель еще загружается. Подождите.")
        result = model.transcribe(file_path, language="ru")
        formatted_text = format_text(result["text"])  # Автоформатирование текста
        return formatted_text
    except Exception as e:
        return f"Ошибка: {e}"

# Функция автоформатирования текста
def format_text(text):
    sentences = text.split(". ")
    formatted_text = ".\n".join(sentences)  # Добавляем перенос строки после точки
    return formatted_text

# Обработчик выбора файла
def select_file():
    file_path = filedialog.askopenfilename(
        title="Выберите аудиофайл",
        filetypes=("Аудиофайлы", "*.mp3 *.wav *.m4a *.ogg"),
    )
    if file_path:
        entry_file_path.delete(0, END)
        entry_file_path.insert(0, file_path)

# Обработчик сохранения в .docx
def save_to_docx(text):
    save_path = filedialog.asksaveasfilename(
        title="Сохраните результат",
        defaultextension=".docx",
        filetypes=("Документ Word", "*.docx"),
    )
    if save_path:
        doc = Document()
        doc.add_heading("Результат транскрибации", level=1)
        doc.add_paragraph(text)
        doc.save(save_path)
        messagebox.showinfo("Успех", "Результат сохранён в файл!")
    else:
        messagebox.showwarning("Внимание", "Файл не был сохранён!")

# Обработчик кнопки "Распознать"
def recognize_audio():
    file_path = entry_file_path.get()
    if not file_path:
        messagebox.showwarning("Ошибка", "Выберите файл для распознавания!")
        return

    btn_recognize.config(state=DISABLED)
    text_result.delete(1.0, END)
    text_result.insert(END, "Распознаю текст...\n")

    def process():
        transcription = transcribe_audio(file_path)
        text_result.delete(1.0, END)
        text_result.insert(END, transcription)
        if transcription and "Ошибка" not in transcription:
            save_to_docx(transcription)
            messagebox.showinfo("Готово", "Транскрибация завершена и сохранена!")
        btn_recognize.config(state=NORMAL)

    Thread(target=process).start()

# Обработчик переключения темы
def change_theme():
    selected_theme = theme_selector.get()
    root.style.theme_use(selected_theme)

# Обработчик для выбора модели с описанием
def select_model(event):
    global model_name
    selected = model_selector.get()
    model_name = selected
    description_label.config(text=MODEL_OPTIONS[selected])

# Обработчик для полноэкранного режима
def toggle_fullscreen():
    root.attributes("-fullscreen", not root.attributes("-fullscreen"))

# Создаём интерфейс
root = ttk.Window(themename="superhero")  # Тема по умолчанию
root.title("Whisper GUI - Распознавание аудио в текст")
root.geometry("800x600")
root.minsize(600, 400)

# Поле для ввода пути к файлу
frame_file = ttk.Frame(root)
frame_file.pack(pady=10, padx=10, fill=X)

label_file = ttk.Label(frame_file, text="Путь к аудиофайлу:", font=("Arial", 12))
label_file.pack(side=LEFT, padx=5)

entry_file_path = ttk.Entry(frame_file, width=50, font=("Arial", 12))
entry_file_path.pack(side=LEFT, padx=5, fill=X, expand=True)

btn_browse = ttk.Button(
    frame_file, text="Обзор", command=select_file, bootstyle="success", padding=5
)
btn_browse.pack(side=LEFT, padx=5)

# Выбор модели
frame_model = ttk.Frame(root)
frame_model.pack(pady=10, padx=10, fill=X)

label_model = ttk.Label(frame_model, text="Выберите модель:", font=("Arial", 12))
label_model.pack(side=LEFT, padx=5)

model_selector = ttk.Combobox(frame_model, values=list(MODEL_OPTIONS.keys()), state="readonly")
model_selector.set("large")  # Модель по умолчанию
model_selector.pack(side=LEFT, padx=5)
model_selector.bind("<<ComboboxSelected>>", select_model)

# Описание модели
description_label = ttk.Label(frame_model, text=MODEL_OPTIONS["large"], wraplength=500)
description_label.pack(side=LEFT, padx=5)

# Кнопка распознавания
btn_recognize = ttk.Button(
    root, text="Распознать", command=recognize_audio, bootstyle="primary", padding=10
)
btn_recognize.pack(pady=10)

# Поле для вывода текста
label_result = ttk.Label(root, text="Результат транскрибации:", font=("Arial", 12))
label_result.pack(anchor=W, padx=10)

text_result = ttk.Text(root, wrap=WORD, font=("Arial", 12))
text_result.pack(pady=5, padx=10, fill=BOTH, expand=True)

# Панель для дополнительных настроек
frame_options = ttk.Frame(root)
frame_options.pack(pady=10, padx=10, fill=X)

# Кнопка для смены темы
theme_selector = ttk.Combobox(
    frame_options, values=root.style.theme_names(), state="readonly"
)
theme_selector.set("superhero")
theme_selector.pack(side=LEFT, padx=5)

btn_theme = ttk.Button(frame_options, text="Применить тему", command=change_theme)
btn_theme.pack(side=LEFT, padx=5)

# Полноэкранный режим
btn_fullscreen = ttk.Button(
    frame_options, text="На весь экран", command=toggle_fullscreen, bootstyle="danger"
)
btn_fullscreen.pack(side=LEFT, padx=5)

# Загружаем модель в фоновом потоке
text_result.insert(END, "Загружаю модель Whisper...\n")
Thread(target=load_model).start()

# Запуск интерфейса
root.mainloop()